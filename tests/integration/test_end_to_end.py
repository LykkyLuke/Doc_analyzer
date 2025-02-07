import pytest
from pathlib import Path
import docx
import json
from src.storage_handler import StorageHandler
from src.document_processor import DocumentProcessor
from src.api_client import GeminiAPIClient
from src.summarization_engine import SummarizationEngine

@pytest.fixture
def test_docx():
    """Create a test .docx file."""
    doc_path = Path("tests/test_data/test_document.docx")
    doc_path.parent.mkdir(parents=True, exist_ok=True)
    
    doc = docx.Document()
    doc.add_paragraph("This is the first paragraph of the test document.")
    doc.add_paragraph("This is the second paragraph with some important information.")
    doc.add_paragraph("This is the third paragraph containing key points.")
    doc.save(doc_path)
    
    yield doc_path
    doc_path.unlink()  # Cleanup

@pytest.fixture
def test_config():
    """Create a test config file."""
    config_path = Path("tests/test_data/config.json")
    config_path.parent.mkdir(parents=True, exist_ok=True)
    
    config = {
        "api_key": "test_api_key",
        "model_settings": {
            "model": "gemini-pro",
            "temperature": 0.7,
            "max_output_tokens": 2048
        }
    }
    
    with open(config_path, 'w') as f:
        json.dump(config, f)
    
    yield config_path
    config_path.unlink()  # Cleanup

def test_document_analysis_workflow(test_docx, test_config):
    """Test the complete document analysis workflow."""
    # Initialize components
    storage = StorageHandler(str(test_config))
    api_client = GeminiAPIClient(storage)
    doc_processor = DocumentProcessor(storage)
    summarizer = SummarizationEngine(api_client, storage)
    
    # Test API key persistence
    storage.save_api_key("test_api_key")
    assert storage.load_api_key() == "test_api_key"
    
    # Test document processing
    document = doc_processor.read_document(str(test_docx))
    text = doc_processor.extract_text(document)
    chunks = doc_processor.create_chunks(text)
    
    assert len(chunks) > 0
    assert all(isinstance(chunk, str) for chunk in chunks)
    assert all(len(chunk) <= doc_processor.chunk_settings['max_chunk_size'] for chunk in chunks)
    
    # Mock API responses for testing
    with pytest.mock.patch.object(api_client, 'generate_content') as mock_generate:
        mock_generate.return_value = "Test summary"
        
        # Test summarization
        summary = summarizer.summarize_document(chunks)
        assert isinstance(summary, str)
        assert len(summary) > 0

def test_error_handling(test_docx, test_config):
    """Test error handling in the workflow."""
    storage = StorageHandler(str(test_config))
    api_client = GeminiAPIClient(storage)
    doc_processor = DocumentProcessor(storage)
    summarizer = SummarizationEngine(api_client, storage)
    
    # Test invalid API key
    with pytest.raises(Exception):
        storage.save_api_key("")
    
    # Test invalid document
    with pytest.raises(FileNotFoundError):
        doc_processor.read_document("nonexistent.docx")
    
    # Test API failure
    with pytest.mock.patch.object(api_client, 'generate_content') as mock_generate:
        mock_generate.side_effect = Exception("API Error")
        with pytest.raises(Exception):
            summarizer.summarize_document(["test chunk"])

def test_progress_tracking(test_docx, test_config):
    """Test progress tracking functionality."""
    storage = StorageHandler(str(test_config))
    api_client = GeminiAPIClient(storage)
    summarizer = SummarizationEngine(api_client, storage)
    
    progress_updates = []
    summarizer.set_progress_callback(
        lambda current, total: progress_updates.append((current, total))
    )
    
    with pytest.mock.patch.object(api_client, 'generate_content') as mock_generate:
        mock_generate.return_value = "Test summary"
        summarizer.summarize_document(["chunk1", "chunk2", "chunk3"])
    
    assert len(progress_updates) == 3
    assert progress_updates[-1] == (3, 3) 