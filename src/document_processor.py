from pathlib import Path
from typing import List
import docx
from docx.document import Document
from logging_config import get_logger
import uuid

class DocumentProcessor:
    """Handles document reading and chunking operations."""
    
    def __init__(self, config_handler):
        """Initialize with configuration handler for chunk settings."""
        self.logger = get_logger(__name__)
        self.logger.info("Initializing DocumentProcessor")
        
        self.config = config_handler
        self.chunk_settings = self.config.load_config().get('chunk_settings', {
            'max_chunk_size': 1000,
            'overlap': 100
        })
        self.logger.debug(f"Chunk settings loaded: {self.chunk_settings}")
    
    def read_document(self, file_path: str) -> Document:
        """Read a .docx file and return the document object."""
        request_id = str(uuid.uuid4())
        self.logger.info(f"Reading document (request_id: {request_id}): {file_path}")
        
        file_path = Path(file_path)
        if not file_path.exists():
            self.logger.error(f"Document not found (request_id: {request_id}): {file_path}")
            raise FileNotFoundError(f"Document not found: {file_path}")
        
        if file_path.suffix.lower() != '.docx':
            self.logger.error(f"Invalid file type (request_id: {request_id}): {file_path.suffix}")
            raise ValueError("Only .docx files are supported")
        
        try:
            document = docx.Document(file_path)
            self.logger.info(f"Document read successfully (request_id: {request_id})")
            self.logger.debug(f"Document paragraphs: {len(document.paragraphs)}")
            return document
        except Exception as e:
            self.logger.error(
                f"Error reading document (request_id: {request_id}): {str(e)}",
                exc_info=True
            )
            raise ValueError(f"Error reading document: {str(e)}")
    
    def extract_text(self, document: Document) -> str:
        """Extract text content from the document."""
        request_id = str(uuid.uuid4())
        self.logger.info(f"Extracting text from document (request_id: {request_id})")
        
        try:
            paragraphs = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
            text = "\n".join(paragraphs)
            
            self.logger.info(f"Text extracted successfully (request_id: {request_id})")
            self.logger.debug(f"Extracted text length: {len(text)} characters")
            self.logger.debug(f"Number of paragraphs: {len(paragraphs)}")
            
            return text
        except Exception as e:
            self.logger.error(
                f"Error extracting text (request_id: {request_id}): {str(e)}",
                exc_info=True
            )
            raise
    
    def create_chunks(self, text: str) -> List[str]:
        """Split text into chunks based on configuration."""
        request_id = str(uuid.uuid4())
        self.logger.info(f"Creating text chunks (request_id: {request_id})")
        self.logger.debug(f"Input text length: {len(text)} characters")
        
        # If text is empty, return empty list
        if not text.strip():
            self.logger.warning(f"Empty text provided (request_id: {request_id})")
            return []
            
        max_size = self.chunk_settings['max_chunk_size']
        overlap = self.chunk_settings['overlap']
        
        try:
            # If text is smaller than max_size, return it as a single chunk
            if len(text) <= max_size:
                self.logger.info(
                    f"Text smaller than max chunk size, returning as single chunk "
                    f"(request_id: {request_id})"
                )
                return [text]
            
            # Split text into sentences (improved implementation)
            sentences = []
            for line in text.split('\n'):
                # Split by common sentence endings
                for sentence in line.replace('!', '.').replace('?', '.').split('.'):
                    sentence = sentence.strip()
                    if sentence:
                        sentences.append(sentence)
            
            self.logger.debug(f"Number of sentences: {len(sentences)}")
            
            # If no sentences found, return text as single chunk
            if not sentences:
                self.logger.warning(
                    f"No sentences found, returning entire text as chunk "
                    f"(request_id: {request_id})"
                )
                return [text]
            
            chunks = []
            current_chunk = []
            current_size = 0
            
            for sentence in sentences:
                sentence_size = len(sentence)
                
                # If a single sentence is larger than max_size, split it
                if sentence_size > max_size:
                    if current_chunk:
                        chunk_text = '. '.join(current_chunk) + '.'
                        chunks.append(chunk_text)
                    
                    # Split long sentence into smaller parts
                    words = sentence.split()
                    current_part = []
                    current_part_size = 0
                    
                    for word in words:
                        if current_part_size + len(word) > max_size:
                            chunks.append(' '.join(current_part))
                            current_part = []
                            current_part_size = 0
                        current_part.append(word)
                        current_part_size += len(word) + 1
                    
                    if current_part:
                        chunks.append(' '.join(current_part))
                    
                    current_chunk = []
                    current_size = 0
                    continue
                
                if current_size + sentence_size > max_size and current_chunk:
                    # Store current chunk
                    chunk_text = '. '.join(current_chunk) + '.'
                    chunks.append(chunk_text)
                    
                    # Keep last few sentences for overlap
                    overlap_sentences = current_chunk[-2:] if len(current_chunk) > 2 else current_chunk
                    current_chunk = overlap_sentences
                    current_size = sum(len(s) for s in current_chunk)
                
                current_chunk.append(sentence)
                current_size += sentence_size
            
            # Add the last chunk if it exists
            if current_chunk:
                chunk_text = '. '.join(current_chunk) + '.'
                chunks.append(chunk_text)
            
            self.logger.info(
                f"Chunking completed (request_id: {request_id}), "
                f"created {len(chunks)} chunks"
            )
            return chunks
            
        except Exception as e:
            self.logger.error(
                f"Error creating chunks (request_id: {request_id}): {str(e)}",
                exc_info=True
            )
            raise
    
    def process_document(self, file_path: str) -> List[str]:
        """Process document and return chunks of text."""
        request_id = str(uuid.uuid4())
        self.logger.info(f"Starting document processing (request_id: {request_id})")
        
        try:
            document = self.read_document(file_path)
            text = self.extract_text(document)
            chunks = self.create_chunks(text)
            
            self.logger.info(
                f"Document processing completed (request_id: {request_id}), "
                f"generated {len(chunks)} chunks"
            )
            return chunks
        except Exception as e:
            self.logger.error(
                f"Document processing failed (request_id: {request_id}): {str(e)}",
                exc_info=True
            )
            raise 